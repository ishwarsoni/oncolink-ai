"""
=============================================================================
generate_samples.py - Create Sample Documents for Testing
=============================================================================

What this script does:
    Creates synthetic oncology documents in different formats (TXT, DOCX, PDF)
    so you can test the OncoLink application without needing real medical files.

How to run:
    python scripts/generate_samples.py

Files created:
    - data/referral.txt              (Text file - already exists)
    - data/lab_report.docx           (Word document)
    - data/pathology_report.pdf      (PDF document)

Python concepts used:
    - Creating files with python-docx
    - Creating files with PyMuPDF
    - os.path for file paths
=============================================================================
"""

# Import required libraries
import os           # For working with file paths
import docx         # For creating Word documents
import fitz         # PyMuPDF - for creating PDFs


def create_docx_sample():
    """
    Create a sample lab report as a DOCX file.
    
    What this does:
        Uses the python-docx library to create a Word document with
        paragraphs of text representing a lab report.
    
    Why DOCX:
        Demonstrates that OncoLink can read Word documents, which doctors
        often use for consultation notes and referral letters.
    """
    
    # Define the output path
    output_path = "data/lab_report.docx"
    
    # Create a new Word document
    document = docx.Document()
    
    # Add a title (bold, larger font)
    document.add_heading("Lab Report", level=1)
    
    # Add patient information section
    document.add_heading("Patient Information", level=2)
    document.add_paragraph("Patient Name: Meera Sharma")
    document.add_paragraph("Date of Birth: 15/06/1966")
    document.add_paragraph("MRN: LAB-2024-0183")
    document.add_paragraph("Date Collected: 2024-01-20")
    
    # Add test results section
    document.add_heading("Complete Blood Count", level=2)
    document.add_paragraph("Hemoglobin: 12.8 g/dL (Reference: 12.0-15.5)")
    document.add_paragraph("WBC: 6.2 x 10^3/uL (Reference: 4.0-10.0)")
    document.add_paragraph("Platelets: 245 x 10^3/uL (Reference: 150-400)")
    
    document.add_heading("Comprehensive Metabolic Panel", level=2)
    document.add_paragraph("All values within normal limits.")
    
    document.add_heading("Tumor Markers", level=2)
    document.add_paragraph("CA 15-3: 32 U/mL (Reference: <30 U/mL)")
    document.add_paragraph("Note: Slightly elevated CA 15-3.")
    
    # Add impression
    document.add_heading("Impression", level=2)
    document.add_paragraph("Normal organ function. Tumor marker slightly elevated, consistent with known breast malignancy.")
    
    # Save the document
    document.save(output_path)
    print(f"  ✅ Created: {output_path}")


def create_pdf_sample():
    """
    Create a sample pathology report as a PDF file.
    
    What this does:
        Uses PyMuPDF (fitz) to create a PDF document with text.
        Each page is created by adding a text block.
    
    Why PDF:
        PDF is the most common format for pathology and radiology reports.
        Demonstrates that OncoLink can extract text from PDFs.
    """
    
    # Define the output path
    output_path = "data/pathology_report.pdf"
    
    # Create a new PDF document
    pdf_document = fitz.open()
    
    # =========================================================================
    # Page 1: Patient info and gross description
    # =========================================================================
    page1 = pdf_document.new_page()
    
    # Define the text content for page 1
    # We use triple-quoted strings for multi-line text
    page1_text = """
CITY PATHOLOGY LAB
    
Pathology Report
    
Patient: Meera Sharma
DOB: 15/06/1966
MRN: PATH-2024-0183
    
Specimen: Left breast core needle biopsy
Date Collected: 2024-01-15
Date Reported: 2024-01-18
    
GROSS DESCRIPTION:
Three core biopsies measuring 1.5 cm, 1.8 cm, and 1.2 cm in length.
All cores submitted for processing.
    
MICROSCOPIC DESCRIPTION:
Sections show invasive ductal carcinoma (IDC), Nottingham histologic
grade 2 (moderately differentiated).
    
Tubule formation score: 2
Nuclear pleomorphism score: 2
Mitotic count score: 1
    
Total Nottingham score: 5/9
"""
    
    # Insert text at position (50, 100) on the page
    # fitz uses points (72 points = 1 inch) for positioning
    page1.insert_text(fitz.Point(50, 100), page1_text, fontsize=10)
    
    # =========================================================================
    # Page 2: Immunohistochemistry and diagnosis
    # =========================================================================
    page2 = pdf_document.new_page()
    
    page2_text = """
IMMUNOHISTOCHEMISTRY RESULTS:
    
ER: Positive (90%, strong intensity)
PR: Positive (70%, moderate intensity)
HER2: Negative (1+ by IHC)
Ki-67: 25%
    
DIAGNOSIS:
Invasive ductal carcinoma of the left breast, Grade 2.
    
Hormone Receptor Status:
    ER Positive
    PR Positive
    HER2 Negative (Triple Negative status: Negative)
    
Pathologic Stage: pT2 (tumor size estimated 3.2 cm on gross)
    
COMMENTS:
All immunohistochemical stains have been reviewed.
Appropriate controls show expected results.
"""
    
    page2.insert_text(fitz.Point(50, 100), page2_text, fontsize=10)
    
    # Save the PDF document
    pdf_document.save(output_path)
    pdf_document.close()
    print(f"  ✅ Created: {output_path}")


def main():
    """
    Main function that creates all sample documents.
    """
    print("Creating sample documents for OncoLink testing...")
    print()
    
    # Create DOCX sample
    create_docx_sample()
    
    # Create PDF sample
    create_pdf_sample()
    
    print()
    print("All sample documents created successfully!")
    print("You can find them in the data/ folder.")


if __name__ == "__main__":
    main()
