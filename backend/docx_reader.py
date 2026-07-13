"""
=============================================================================
docx_reader.py - Extract text from DOCX files
=============================================================================

What this module does:
    Takes a DOCX file (Word document) as raw bytes and extracts all text.
    Uses the python-docx library to read the document.

What is a DOCX file?
    A DOCX file is actually a ZIP file containing XML documents.
    When you "open" a .docx file, Word reads the XML and renders it.
    The python-docx library handles all this complexity for us.

Why a separate module:
    Same reason as pdf_reader.py - encapsulation.
    The rest of our application doesn't need to know how DOCX reading works.
    It just calls read_docx() and gets text back.
=============================================================================
"""

# Import the python-docx library
import docx

# Import io (input/output) for working with bytes in memory
# BytesIO creates a "file-like object" from bytes.
# Some libraries expect a file path or file object, not raw bytes.
# BytesIO wraps our bytes so it looks like a file to those libraries.
import io


def read_docx(file_bytes, filename):
    """
    Extract all text from a DOCX (Word) file.
    
    Parameters:
        file_bytes (bytes): The raw bytes of the uploaded DOCX file.
            Same concept as pdf_reader.py - the file content in memory.
        
        filename (str): The name of the file (for error messages).
    
    Returns:
        str: All paragraphs from the DOCX, joined with newlines.
            If an error occurs, returns an error message string.
    
    How DOCX files store text:
        A DOCX file organizes text into "paragraphs".
        Each paragraph is a block of text (like a paragraph in a book).
        Unlike PDFs, there are no "pages" in the DOCX data model -
        pages are created by Word when you open the file.
    
    Python concepts used:
        - Importing modules (docx, io)
        - Creating objects (io.BytesIO, docx.Document)
        - For loop over a collection (document.paragraphs)
        - Accessing object attributes (paragraph.text)
        - try/except error handling
    """
    
    try:
        # =========================================================================
        # Step 1: Wrap the bytes in a file-like object
        # =========================================================================
        # python-docx's Document() can open a file from disk (give it a path)
        # or from a "file-like object" in memory.
        # 
        # io.BytesIO() takes our raw bytes and creates an object that behaves
        # like an open file. This lets us work with the document without saving
        # it to disk first.
        #
        # Think of BytesIO as a "pretend file" that lives in RAM.
        file_like_object = io.BytesIO(file_bytes)
        
        # =========================================================================
        # Step 2: Open the DOCX document
        # =========================================================================
        # docx.Document() reads the DOCX file and parses all its content.
        # It creates a Document object that we can ask for paragraphs, tables,
        # headers, etc.
        document = docx.Document(file_like_object)
        
        # Initialize an empty string for all the text
        all_text = ""
        
        # =========================================================================
        # Step 3: Extract text from each paragraph
        # =========================================================================
        # document.paragraphs is a list of all paragraphs in the document.
        # Each paragraph object has a .text attribute containing the text.
        # 
        # We loop through every paragraph and add its text to our string.
        # We add a newline (\n) after each paragraph to preserve the document's
        # paragraph structure.
        for paragraph in document.paragraphs:
            all_text = all_text + paragraph.text + "\n"
        
        return all_text
    
    except Exception as error:
        # Handle any errors (corrupted file, wrong format, etc.)
        error_message = f"Error reading DOCX file '{filename}': {error}"
        return error_message
